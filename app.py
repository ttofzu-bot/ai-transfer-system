"""
AI Transfer System V5.0 — FZÚ Patent & Research Intelligence
=============================================================
New: patent statistics, XLSX table export, fixed DOCX (no **), transparent relevance, country extraction
Pipeline: PDF → Gemini → Google Patents → AI filtr → OpenAlex → Stats → Analýza → Docs + Excel
"""

import streamlit as st
import requests
import json
import time
import os
import io
import re
import random
from datetime import datetime
from pathlib import Path
from collections import Counter
import pandas as pd

st.set_page_config(page_title="AI Transfer System — FZÚ", page_icon="🔬", layout="wide", initial_sidebar_state="expanded")
HISTORY_DIR = Path("analysis_history")
HISTORY_DIR.mkdir(exist_ok=True)

# ---------------------------------------------------------------------------
# CSS (same as V4 with fixes)
# ---------------------------------------------------------------------------
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Instrument+Sans:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');
html, body, [class*="st-"] { font-family: 'Instrument Sans', -apple-system, BlinkMacSystemFont, sans-serif !important; }
.block-container { max-width: 1100px; padding: 2.5rem 1.5rem 4rem; }
#MainMenu, footer { visibility: hidden; }
[data-testid="stSidebar"] { background: #f7f9fc !important; border-right: 1px solid #e8ecf1 !important; }
[data-testid="stSidebar"] h3 { color: #0f172a !important; font-size: 0.8rem !important; text-transform: uppercase; letter-spacing: 0.08em !important; font-weight: 600 !important; }
[data-testid="stSidebar"] .stTextInput input { background: #fff !important; border: 1px solid #d1d9e6 !important; border-radius: 8px !important; }
[data-testid="stSidebar"] button { background: #fff !important; border: 1px solid #d1d9e6 !important; border-radius: 8px !important; width: 100%; font-weight: 500 !important; }
[data-testid="stSidebar"] button:hover { background: #f0f9ff !important; border-color: #0ea5e9 !important; }
.hero { background: linear-gradient(135deg, #0b1120 0%, #162240 40%, #1e3a5f 70%, #0b1120 100%); color: white; padding: 3rem 3rem 2.5rem; border-radius: 20px; margin-bottom: 2rem; position: relative; overflow: hidden; border: 1px solid rgba(56,189,248,0.1); }
.hero-badge { display: inline-block; background: rgba(56,189,248,0.12); color: #7dd3fc; padding: 4px 14px; border-radius: 20px; font-size: 0.7rem; font-weight: 600; letter-spacing: 0.06em; text-transform: uppercase; margin-bottom: 1rem; border: 1px solid rgba(56,189,248,0.15); }
.hero h1 { font-size: 2rem; font-weight: 700; margin: 0 0 0.6rem; letter-spacing: -0.03em; background: linear-gradient(135deg, #fff 0%, #94cef5 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
.hero p { font-size: 0.92rem; opacity: 0.55; margin: 0; }
.hero-org { font-size: 0.75rem; opacity: 0.35; margin-top: 1rem; font-weight: 500; }
.phase-bar { display: flex; gap: 4px; margin: 1.5rem 0; flex-wrap: wrap; }
.phase-pill { padding: 7px 16px; border-radius: 8px; font-size: 0.72rem; font-weight: 600; border: 1px solid transparent; }
.phase-active { background: #0c4a6e; color: #7dd3fc; border-color: rgba(56,189,248,0.3); }
.phase-done { background: #064e3b; color: #6ee7b7; border-color: rgba(16,185,129,0.3); }
.phase-pending { background: #f1f5f9; color: #94a3b8; }
.stat-row { display: flex; gap: 14px; margin: 1.5rem 0; }
.stat-box { flex: 1; background: #fff; border: 1px solid #e8ecf1; border-radius: 14px; padding: 1.25rem 1rem; text-align: center; }
.stat-box .num { font-size: 1.6rem; font-weight: 700; color: #0f172a; }
.stat-box .label { font-size: 0.7rem; color: #64748b; margin-top: 4px; font-weight: 500; text-transform: uppercase; letter-spacing: 0.04em; }
.patent-card { border: 1px solid #eaeff5; border-radius: 14px; padding: 1.25rem 1.5rem; margin-bottom: 10px; background: #fff; }
.patent-card:hover { border-color: #93c5fd; }
.patent-card h4 { font-size: 0.9rem; font-weight: 600; margin: 0 0 0.4rem; color: #0f172a; }
.patent-card .meta { font-size: 0.78rem; color: #64748b; line-height: 1.7; }
.patent-card .applicant { display: inline-block; background: #eff6ff; color: #1d4ed8; padding: 3px 10px; border-radius: 6px; font-size: 0.7rem; font-weight: 600; margin-top: 0.5rem; margin-right: 4px; }
.relevance { display: inline-block; padding: 3px 10px; border-radius: 6px; font-size: 0.7rem; font-weight: 600; margin-top: 0.5rem; }
.rel-high { background: #ecfdf5; color: #065f46; border: 1px solid #a7f3d0; }
.rel-med { background: #fefce8; color: #854d0e; border: 1px solid #fde68a; }
.rel-low { background: #fef2f2; color: #991b1b; border: 1px solid #fecaca; }
.analysis-box { background: #fff; border: 1px solid #e2e8f0; border-left: 4px solid #0ea5e9; border-radius: 0 14px 14px 0; padding: 2rem; margin: 1rem 0; font-size: 0.88rem; line-height: 1.8; }
.stButton > button[kind="primary"] { background: linear-gradient(135deg, #0369a1, #0ea5e9) !important; color: white !important; border: none !important; border-radius: 10px !important; font-weight: 600 !important; }
.stTabs [aria-selected="true"] { background: #f0f9ff !important; color: #0369a1 !important; }
[data-testid="stFileUploader"] { border: 2px dashed #d1d9e6 !important; border-radius: 14px !important; background: #f8fafc !important; }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# SESSION STATE
# ---------------------------------------------------------------------------
defaults = {
    "phase": 0, "pdf_text": "", "pdf_name": "",
    "tech_summary": "", "tech_keywords_en": "", "tech_domains": "",
    "search_queries": [],
    "patents_raw": [], "patents_filtered": [],
    "openalex_results": [], "analysis": "", "doc_content": "", "xlsx_content": "",
    "page": "main",
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ---------------------------------------------------------------------------
# SIDEBAR
# ---------------------------------------------------------------------------
with st.sidebar:
    st.markdown("### ⚙️ Nastavení")
    st.caption("API klíče platí jen pro tuto session.")
    gemini_key = st.text_input("Gemini API Key", value=os.environ.get("GEMINI_API_KEY", ""), type="password")
    serpapi_key = st.text_input("SerpApi Key", value=os.environ.get("SERPAPI_KEY", ""), type="password")
    st.divider()
    st.markdown("### 📊 Parametry")
    max_patents_per_query = st.slider("Patentů na dotaz", 10, 50, 25)
    max_openalex = st.slider("Max článků z OpenAlex", 5, 50, 30)
    relevance_threshold = st.slider("Min. relevance (0-10)", 0, 10, 5)
    st.divider()
    if st.button("📂 Historie analýz"): st.session_state.page = "history"; st.rerun()
    if st.button("🔬 Nová analýza"): st.session_state.page = "main"; st.rerun()
    st.divider()
    st.caption("AI Transfer System V5.0\nFZÚ AV ČR")

# ---------------------------------------------------------------------------
# GEMINI WITH MULTI-MODEL FALLBACK
# ---------------------------------------------------------------------------
def call_gemini(api_key, prompt, system_instruction="", max_retries=5):
    models = ["gemini-2.5-flash", "gemini-2.5-flash-lite", "gemini-1.5-flash"]
    for model in models:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
        payload = {"contents": [{"parts": [{"text": prompt}]}]}
        if system_instruction:
            payload["systemInstruction"] = {"parts": [{"text": system_instruction}]}
        payload["generationConfig"] = {"temperature": 0.2, "maxOutputTokens": 8192}
        for attempt in range(max_retries):
            try:
                resp = requests.post(url, json=payload, timeout=120)
                if resp.status_code == 200:
                    return resp.json()["candidates"][0]["content"]["parts"][0]["text"]
                elif resp.status_code in (429, 503):
                    wait = (2 ** attempt) * 2 + random.uniform(1, 3)
                    if attempt < max_retries - 1:
                        st.toast(f"⏳ {model} přetížen, čekám {wait:.0f}s...")
                        time.sleep(wait)
                    else:
                        break
                else:
                    raise Exception(f"Gemini {resp.status_code}: {resp.text[:300]}")
            except requests.exceptions.Timeout:
                if attempt < max_retries - 1: time.sleep(3)
                else: break
        st.toast(f"⚠️ {model} nedostupný, zkouším další...")
        time.sleep(2)
    raise Exception("Všechny Gemini modely přetížené. Zkus za chvíli.")

# ---------------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------------
def extract_pdf_text(f):
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(f)
        return "\n".join(p.extract_text() for p in reader.pages if p.extract_text())
    except Exception as e:
        st.error(f"Chyba při čtení PDF: {e}")
        return ""

def extract_country(patent_id):
    """Extract filing country from patent ID like US20210001234A1 → US"""
    if not patent_id or patent_id == "—": return "—"
    m = re.match(r'^([A-Z]{2})', patent_id)
    return m.group(1) if m else "—"

def extract_year(date_str):
    """Extract year from various date formats."""
    if not date_str or date_str == "—": return None
    m = re.search(r'(\d{4})', str(date_str))
    return int(m.group(1)) if m else None

def strip_markdown(text):
    """Remove ** and other markdown from text for clean docx output."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # **bold** → bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)  # *italic* → italic
    text = re.sub(r'#{1,6}\s*', '', text)  # ### heading → heading
    text = re.sub(r'`(.*?)`', r'\1', text)  # `code` → code
    return text

def compute_stats(patents):
    """Compute patent statistics: top assignees, filings by year."""
    # Top assignees
    assignees = [p["applicant"] for p in patents if p["applicant"] != "—"]
    assignee_counts = Counter(assignees).most_common(20)

    # Filings by year
    years = [extract_year(p["filing_date"]) for p in patents]
    years = [y for y in years if y and 2000 <= y <= 2030]
    year_counts = Counter(years)
    year_range = range(min(year_counts.keys()) if year_counts else 2000,
                       (max(year_counts.keys()) if year_counts else 2025) + 1)
    yearly = [(y, year_counts.get(y, 0)) for y in year_range]

    # Countries
    countries = [extract_country(p["pub_number"]) for p in patents]
    country_counts = Counter(c for c in countries if c != "—").most_common(10)

    return {"assignees": assignee_counts, "yearly": yearly, "countries": country_counts}

# ---------------------------------------------------------------------------
# STEP 1: ANALYZE DOCUMENT
# ---------------------------------------------------------------------------
def analyze_document(api_key, doc_text):
    system = """You are an expert in patent research and technology transfer.

TASK: Based on the technical document, generate:
1. A concise technology summary in Czech (2-3 sentences)
2. A LIST of ALL potential application domains mentioned or implied in the document (in English)
3. 5-8 English keywords/keyphrases for searching scientific databases
4. FIVE different Google Patents search queries — each targeting a DIFFERENT application domain or angle:
   - QUERY 1 (CORE TECHNOLOGY): The exact technology/method/synthesis described
   - QUERY 2 (APPLICATION 1): First application domain (e.g. electronics, sensors)
   - QUERY 3 (APPLICATION 2): Second application domain (e.g. biomedicine, coatings)
   - QUERY 4 (APPLICATION 3): Third application domain (e.g. energy, catalysis, optics)
   - QUERY 5 (MATERIAL): The base material in broader contexts

CRITICAL RULES:
- You MUST cover ALL application areas mentioned in the document, not just one
- If the document mentions electronics, sensors, biomedicine, shielding, catalysis — generate queries for ALL of them
- Do NOT focus on just one property or application
- Each query max 2-3 phrases with AND/OR, in quotes, English
- NO markdown, NO explanations

OUTPUT FORMAT (follow exactly):
SUMMARY: [Czech summary]
DOMAINS: [comma-separated list of application domains in English]
KEYWORDS: [comma-separated English keywords]
QUERY1: [core technology query]
QUERY2: [application 1 query]
QUERY3: [application 2 query]
QUERY4: [application 3 query]
QUERY5: [material query]"""
    result = call_gemini(api_key, f"Analyze this document thoroughly — identify ALL applications and domains:\n\n{doc_text[:8000]}", system)
    summary, keywords, domains, queries = "", "", "", []
    for line in result.strip().split("\n"):
        line = line.strip()
        if line.startswith("SUMMARY:"): summary = line[8:].strip()
        elif line.startswith("KEYWORDS:"): keywords = line[9:].strip()
        elif line.startswith("DOMAINS:"): domains = line[8:].strip()
        elif line.startswith("QUERY"):
            q = line.split(":", 1)[-1].strip().strip("`\"'")
            q = re.sub(r"^```.*\n?", "", q); q = re.sub(r"\n?```$", "", q)
            if q: queries.append(q.strip())
    return {"summary": summary or "—", "keywords": keywords or "—", "domains": domains or "—", "queries": queries or [result.strip()]}

# ---------------------------------------------------------------------------
# STEP 2: SEARCH GOOGLE PATENTS
# ---------------------------------------------------------------------------
def sanitize_patent_query(query):
    """Clean up a query string so Google Patents / SerpApi can process it."""
    q = query.strip()
    # Remove markdown code block wrappers only
    q = re.sub(r'^```[a-z]*\s*', '', q)
    q = re.sub(r'\s*```$', '', q)
    # Remove wrapping backticks only (not internal quotes)
    if q.startswith('`') and q.endswith('`'):
        q = q[1:-1]
    # Remove field-specific prefixes that SerpApi doesn't support
    q = re.sub(r'\b(title|abstract|claim|applicant|inventor|classification):', '', q, flags=re.IGNORECASE)
    # Remove stray backslashes and brackets
    q = q.replace('\\', '').replace('{', '').replace('}', '').replace('[', '').replace(']', '')
    # Normalize whitespace
    q = re.sub(r'\s+', ' ', q).strip()
    # Wrap in parentheses if not already
    if q and not q.startswith('('):
        q = f"({q})"
    return q

def search_google_patents(api_key, query, max_results=25):
    query = sanitize_patent_query(query)
    patents, page_num = [], 1
    while len(patents) < max_results:
        params = {"engine": "google_patents", "q": query, "api_key": api_key, "num": 100, "page": page_num}
        resp = requests.get("https://serpapi.com/search.json", params=params, timeout=30)
        if resp.status_code != 200:
            raise Exception(f"SerpApi {resp.status_code}: {resp.text[:200]}")
        data = resp.json()
        # Check for API error
        if "error" in data:
            raise Exception(f"SerpApi error: {data['error']}")
        results = data.get("organic_results", [])
        if not results:
            break
        for r in results:
            if len(patents) >= max_results: break
            if r.get("is_scholar"): continue
            pid = r.get("patent_id", "—") or "—"
            patents.append({
                "title": r.get("title", "—") or "—",
                "abstract": (r.get("snippet", "—") or "—")[:500],
                "applicant": r.get("assignee", "—") or "—",
                "inventor": r.get("inventor", "—") or "—",
                "pub_number": pid,
                "filing_date": r.get("filing_date", "—") or "—",
                "grant_date": r.get("grant_date", "—") or "—",
                "country": extract_country(pid),
                "cpc": r.get("cpc", "—") or "—",
                "pdf_link": r.get("pdf", ""),
                "gp_link": r.get("patent_link", r.get("link", "")),
                "relevance": 0, "rel_type": "—", "rel_reason": "Nefiltrováno",
            })
        if not data.get("organic_results") or not data.get("serpapi_pagination", {}).get("next"): break
        page_num += 1; time.sleep(0.3)
    return patents

# ---------------------------------------------------------------------------
# STEP 3: AI RELEVANCE FILTER
# ---------------------------------------------------------------------------
def filter_patents(api_key, patents, tech_summary, threshold=5):
    if not patents: return []
    system = """You are a patent relevance expert. Score each patent vs the given technology.
For EACH patent return one line: NUMBER|SCORE|TYPE|REASON
- SCORE 0-10 (10=identical technology, 7-9=very relevant, 4-6=partial, 0-3=irrelevant)
- TYPE = COMPETITOR/PARTNER/CUSTOMER/IRRELEVANT
- REASON = max 10 words
EVALUATION CRITERIA (be transparent):
- Compare the patent ABSTRACT with the technology description
- COMPETITOR: similar synthesis/manufacturing method OR identical product
- CUSTOMER: product/device that could directly use this technology as input
- PARTNER: works in closely related field, complementary capability
- IRRELEVANT: keyword overlap only, different domain entirely
BE STRICT. Only NUMBER|SCORE|TYPE|REASON lines."""

    BATCH = 15
    scores = {}
    for start in range(0, len(patents), BATCH):
        batch = patents[start:start+BATCH]
        plist = "\n".join(f"{start+i}. {p['title']} | {p['applicant']} | {p['abstract'][:150]}" for i, p in enumerate(batch))
        try:
            result = call_gemini(api_key, f"TECHNOLOGY:\n{tech_summary}\n\nPATENTS:\n{plist}", system)
            for line in result.strip().split("\n"):
                parts = line.strip().split("|")
                if len(parts) >= 3:
                    try:
                        idx = int(parts[0].strip().rstrip("."))
                        scores[idx] = (min(int(parts[1].strip()), 10), parts[2].strip(), parts[3].strip() if len(parts)>3 else "—")
                    except (ValueError, IndexError):
                        continue
        except Exception as e:
            st.toast(f"⚠️ Filtr batch {start//BATCH + 1} selhal: {str(e)[:50]}")
        time.sleep(0.5)

    filtered = []
    for i, p in enumerate(patents):
        if i in scores: p["relevance"], p["rel_type"], p["rel_reason"] = scores[i]
        else: p["relevance"], p["rel_type"], p["rel_reason"] = 2, "—", "Neohodnoceno"
        if p["relevance"] >= threshold: filtered.append(p)
    filtered.sort(key=lambda x: x["relevance"], reverse=True)
    return filtered

# ---------------------------------------------------------------------------
# STEP 4: OPENALEX
# ---------------------------------------------------------------------------
def search_openalex(keywords, max_results=30):
    query = re.sub(r'["\(\)]', '', keywords.strip())
    parts = [k.strip() for k in query.split(",")][:5]
    query = " ".join(parts)
    if not query: return []
    resp = requests.get("https://api.openalex.org/works",
        params={"search": query, "per_page": max_results, "sort": "relevance_score:desc", "mailto": "transfer@fzu.cz"}, timeout=30)
    if resp.status_code != 200: return []
    results = []
    for w in resp.json().get("results", []):
        insts, is_com = [], False
        for a in w.get("authorships", []):
            for inst in a.get("institutions", []):
                n = inst.get("display_name", "")
                insts.append(n)
                if inst.get("type") in ("company","facility") or any(t in n.lower() for t in ["inc.","ltd.","gmbh","a.s.","corp.","co.","llc","ag"]): is_com = True
        results.append({"title": w.get("title","—") or "—", "year": w.get("publication_year","—"),
            "cited_by": w.get("cited_by_count",0), "institutions": list(set(insts))[:5], "is_commercial": is_com,
            "doi": w.get("doi",""), "type": w.get("type","")})
    return results

# ---------------------------------------------------------------------------
# STEP 5: ANALYSIS
# ---------------------------------------------------------------------------
def run_analysis(api_key, tech_summary, patents, openalex, pdf_text, stats):
    system = """You are a CRITICAL technology transfer expert at a Czech TTO.
INSTRUCTIONS:
- Be REALISTIC, not optimistic
- If data shows weak signals, say it honestly
- Do not suggest partners without evidence
- Write in CZECH
- Do NOT use markdown formatting (no **, no ##, no *, no `)
- Use plain text only. For emphasis, use CAPS or write "Důležité:" before a sentence."""

    assignee_table = "\n".join(f"  {name}: {count} patentů" for name, count in stats["assignees"][:10])
    yearly_str = "\n".join(f"  {y}: {c} patentů" for y, c in stats["yearly"] if c > 0)

    pat_sum = "\n".join(f"- {p['title']} | {p['applicant']} | {p.get('relevance','?')}/10 | {p.get('rel_type','?')} | {p.get('rel_reason','')}" for p in patents[:20])
    pub_sum = "\n".join(f"- {p['title']} ({p['year']}) | Citací: {p['cited_by']} | {', '.join(p['institutions'][:2])}" for p in openalex[:15])
    commercial = [r for r in openalex if r["is_commercial"]]
    com_sum = "\n".join(f"- {c['title']} | {', '.join(c['institutions'][:3])}" for c in commercial) if commercial else "Žádné."

    prompt = f"""TECHNOLOGIE: {tech_summary}

TOP PŘIHLAŠOVATELÉ PATENTŮ (dle četnosti):
{assignee_table}

VÝVOJ PATENTOVÉ AKTIVITY PO LETECH:
{yearly_str}

RELEVANTNÍ PATENTY ({len(patents)}):
{pat_sum}

PUBLIKACE ({len(openalex)}):
{pub_sum}

KOMERČNÍ SPOLUPRÁCE V PUBLIKACÍCH:
{com_sum}

ZDROJOVÝ DOKUMENT:
{pdf_text[:2000]}

Proveď KRITICKOU analýzu v češtině (BEZ markdown formátování, BEZ hvězdiček):
1. SHRNUTÍ TECHNOLOGIE (2-3 věty)
2. PATENTOVÁ KRAJINA
   - Hlavní hráči (na základě tabulky přihlašovatelů)
   - Trend patentové aktivity (roste/klesá/stagnuje)
   - Geografické rozložení
3. KOMERČNÍ SIGNÁLY (konkrétní důkazy z dat)
4. POTENCIÁLNÍ PARTNEŘI (pouze s důkazem v datech)
5. RIZIKA A SLABINY
6. GO / CONDITIONAL GO / NO-GO
7. DOPORUČENÝ DALŠÍ POSTUP"""

    return call_gemini(api_key, prompt, system)

# ---------------------------------------------------------------------------
# XLSX EXPORT
# ---------------------------------------------------------------------------
def generate_xlsx(patents, openalex, tech_summary):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()

    # Sheet 1: Patents
    ws = wb.active
    ws.title = "Patenty"
    headers = ["#", "Název", "Vynálezce", "Majitel", "Datum podání", "Stát", "Číslo patentu", "Klasifikace (CPC)", "Relevance", "Typ", "Důvod relevance", "Abstrakt", "Odkaz"]
    hfont = Font(bold=True, size=10, color="FFFFFF")
    hfill = PatternFill(start_color="0C4A6E", end_color="0C4A6E", fill_type="solid")
    thin = Side(style='thin', color='D1D9E6')
    border = Border(bottom=thin)

    for col, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = hfont; c.fill = hfill; c.alignment = Alignment(horizontal='center', wrap_text=True)

    for idx, p in enumerate(patents, 1):
        row = idx + 1
        ws.cell(row=row, column=1, value=idx)
        ws.cell(row=row, column=2, value=p["title"])
        ws.cell(row=row, column=3, value=p.get("inventor", "—"))
        ws.cell(row=row, column=4, value=p["applicant"])
        ws.cell(row=row, column=5, value=p.get("filing_date", "—"))
        ws.cell(row=row, column=6, value=p.get("country", extract_country(p.get("pub_number",""))))
        ws.cell(row=row, column=7, value=p["pub_number"])
        ws.cell(row=row, column=8, value=p.get("cpc", "—"))
        ws.cell(row=row, column=9, value=p.get("relevance", 0))
        ws.cell(row=row, column=10, value=p.get("rel_type", "—"))
        ws.cell(row=row, column=11, value=p.get("rel_reason", "—"))
        ws.cell(row=row, column=12, value=p.get("abstract", "—"))
        ws.cell(row=row, column=13, value=p.get("gp_link", ""))
        for col in range(1, 14):
            ws.cell(row=row, column=col).font = Font(size=9)
            ws.cell(row=row, column=col).alignment = Alignment(wrap_text=True, vertical='top')
            ws.cell(row=row, column=col).border = border

    ws.column_dimensions['A'].width = 4
    ws.column_dimensions['B'].width = 40
    ws.column_dimensions['C'].width = 25
    ws.column_dimensions['D'].width = 25
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 6
    ws.column_dimensions['G'].width = 18
    ws.column_dimensions['H'].width = 20
    ws.column_dimensions['I'].width = 8
    ws.column_dimensions['J'].width = 12
    ws.column_dimensions['K'].width = 25
    ws.column_dimensions['L'].width = 50
    ws.column_dimensions['M'].width = 30
    ws.auto_filter.ref = ws.dimensions

    # Sheet 2: Publications
    ws2 = wb.create_sheet("Publikace")
    pub_headers = ["#", "Název", "Rok", "Citací", "Komerční", "Instituce", "Typ", "DOI"]
    for col, h in enumerate(pub_headers, 1):
        c = ws2.cell(row=1, column=col, value=h)
        c.font = hfont; c.fill = hfill
    for idx, p in enumerate(openalex, 1):
        row = idx + 1
        ws2.cell(row=row, column=1, value=idx)
        ws2.cell(row=row, column=2, value=p["title"])
        ws2.cell(row=row, column=3, value=p["year"])
        ws2.cell(row=row, column=4, value=p["cited_by"])
        ws2.cell(row=row, column=5, value="Ano" if p["is_commercial"] else "Ne")
        ws2.cell(row=row, column=6, value=", ".join(p["institutions"][:4]))
        ws2.cell(row=row, column=7, value=p["type"])
        ws2.cell(row=row, column=8, value=p.get("doi", ""))
    ws2.column_dimensions['B'].width = 50
    ws2.column_dimensions['F'].width = 40
    ws2.auto_filter.ref = ws2.dimensions

    # Sheet 3: Statistics
    ws3 = wb.create_sheet("Statistiky")
    ws3.cell(row=1, column=1, value="Top přihlašovatelé patentů").font = Font(bold=True, size=12)
    ws3.cell(row=2, column=1, value="Firma/Instituce").font = Font(bold=True)
    ws3.cell(row=2, column=2, value="Počet patentů").font = Font(bold=True)
    stats = compute_stats(patents)
    for i, (name, count) in enumerate(stats["assignees"], 3):
        ws3.cell(row=i, column=1, value=name)
        ws3.cell(row=i, column=2, value=count)
    ws3.column_dimensions['A'].width = 40
    ws3.column_dimensions['B'].width = 15

    row_start = len(stats["assignees"]) + 5
    ws3.cell(row=row_start, column=1, value="Patentová aktivita po letech").font = Font(bold=True, size=12)
    ws3.cell(row=row_start+1, column=1, value="Rok").font = Font(bold=True)
    ws3.cell(row=row_start+1, column=2, value="Počet patentů").font = Font(bold=True)
    for i, (year, count) in enumerate(stats["yearly"], row_start+2):
        ws3.cell(row=i, column=1, value=year)
        ws3.cell(row=i, column=2, value=count)

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf

# ---------------------------------------------------------------------------
# DOCX EXPORT (fixed — no ** markdown)
# ---------------------------------------------------------------------------
def generate_docx(tech_summary, queries, patents_raw_count, patents, openalex, analysis, pdf_filename, threshold, stats):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(10.5)
    s.paragraph_format.space_after = Pt(6)

    # Title
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI Transfer System — Rešeršní zpráva"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(15, 23, 42)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Zdroj: {pdf_filename}  |  {datetime.now().strftime('%d. %m. %Y %H:%M')}"); r.font.size = Pt(10); r.font.color.rgb = RGBColor(100, 116, 139)
    doc.add_page_break()

    # 1. Summary
    doc.add_heading("1. Shrnutí technologie", level=1)
    doc.add_paragraph(strip_markdown(tech_summary))

    # 2. Queries
    doc.add_heading("2. Vyhledávací dotazy", level=1)
    for i, q in enumerate(queries, 1):
        p = doc.add_paragraph(); r = p.add_run(f"Dotaz {i}: "); r.bold = True
        r2 = p.add_run(q); r2.font.name = "Consolas"; r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(3, 105, 161)

    # 3. Assignee frequency table
    doc.add_heading("3. Přehled přihlašovatelů patentů", level=1)
    doc.add_paragraph(f"Celkem nalezeno {patents_raw_count} patentů, po filtraci relevance: {len(patents)}")
    if stats["assignees"]:
        t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
        for i, h in enumerate(["#", "Firma / Instituce", "Počet patentů"]):
            c = t.rows[0].cells[i]; c.text = h
            for pg in c.paragraphs:
                for rn in pg.runs: rn.bold = True; rn.font.size = Pt(9)
        for idx, (name, count) in enumerate(stats["assignees"][:15], 1):
            row = t.add_row().cells
            row[0].text = str(idx); row[1].text = name; row[2].text = str(count)
            for c in row:
                for pg in c.paragraphs:
                    for rn in pg.runs: rn.font.size = Pt(9)

    # 4. Yearly filing trend
    doc.add_heading("4. Vývoj patentové aktivity po letech", level=1)
    if stats["yearly"]:
        t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
        t.rows[0].cells[0].text = "Rok"; t.rows[0].cells[1].text = "Počet patentů"
        for pg in t.rows[0].cells[0].paragraphs:
            for rn in pg.runs: rn.bold = True; rn.font.size = Pt(9)
        for pg in t.rows[0].cells[1].paragraphs:
            for rn in pg.runs: rn.bold = True; rn.font.size = Pt(9)
        for year, count in stats["yearly"]:
            if count > 0:
                row = t.add_row().cells
                row[0].text = str(year); row[1].text = str(count)

    # 5. Patent table
    doc.add_heading("5. Tabulka patentů (filtrované)", level=1)
    if patents:
        t = doc.add_table(rows=1, cols=7); t.style = "Light Grid Accent 1"
        for i, h in enumerate(["#", "Název", "Majitel", "Vynálezce", "Stát", "Datum", "Rel."]):
            c = t.rows[0].cells[i]; c.text = h
            for pg in c.paragraphs:
                for rn in pg.runs: rn.bold = True; rn.font.size = Pt(8)
        for idx, pat in enumerate(patents[:30], 1):
            row = t.add_row().cells
            row[0].text = str(idx); row[1].text = pat["title"][:60]; row[2].text = pat["applicant"][:30]
            row[3].text = pat.get("inventor","—")[:30]; row[4].text = pat.get("country","—")
            row[5].text = pat.get("filing_date","—"); row[6].text = f"{pat.get('relevance','?')}/10"
            for c in row:
                for pg in c.paragraphs:
                    for rn in pg.runs: rn.font.size = Pt(8)

    # 6. Publications
    doc.add_heading("6. Vědecké publikace (OpenAlex)", level=1)
    doc.add_paragraph(f"Celkem: {len(openalex)} článků")
    commercial = [r for r in openalex if r["is_commercial"]]
    if commercial:
        doc.add_heading("Články s komerční spoluprací:", level=2)
        for r in commercial[:10]:
            p = doc.add_paragraph(style="List Bullet")
            rn = p.add_run(f"{r['title']} ({r['year']})"); rn.bold = True; rn.font.size = Pt(9.5)
            p.add_run(f"\n   Instituce: {', '.join(r['institutions'][:3])}")

    # 7. AI Analysis (strip markdown!)
    doc.add_heading("7. AI Analýza komerčního potenciálu", level=1)
    clean_analysis = strip_markdown(analysis)
    for line in clean_analysis.split("\n"):
        line = line.strip()
        if not line: continue
        # Detect section headers (lines ending with :)
        if re.match(r'^\d+\.', line) or line.endswith(':'):
            p = doc.add_paragraph()
            r = p.add_run(line); r.bold = True; r.font.size = Pt(11)
        else:
            doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("Metodologie a nastavení", level=1)
    doc.add_paragraph(f"Systém: AI Transfer System V5.0")
    doc.add_paragraph(f"Datum: {datetime.now().strftime('%d. %m. %Y %H:%M')}")
    doc.add_paragraph(f"Zdrojový dokument: {pdf_filename}")
    doc.add_paragraph(f"Patentová databáze: Google Patents (via SerpApi)")
    doc.add_paragraph(f"Vědecké publikace: OpenAlex")
    doc.add_paragraph(f"AI model: Google Gemini (2.5 Flash / 2.5 Flash-Lite / 1.5 Flash fallback)")
    doc.add_paragraph(f"Práh relevance: {threshold}/10")
    doc.add_paragraph(f"Počet patentů na dotaz: max {len(patents)} (po filtraci z {patents_raw_count})")
    doc.add_paragraph(f"Použité vyhledávací dotazy:")
    for i, q in enumerate(queries, 1):
        p = doc.add_paragraph(f"  Dotaz {i}: {q}")
    doc.add_paragraph("")
    doc.add_paragraph("Výsledky vyžadují odbornou validaci pracovníkem TTO. AI analýza může obsahovat nepřesnosti.")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ---------------------------------------------------------------------------
# HISTORY
# ---------------------------------------------------------------------------
def save_to_history(data):
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    name = re.sub(r'[^\w]', '_', data.get("pdf_name", "unknown"))[:40]
    with open(HISTORY_DIR / f"{ts}_{name}.json", "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2, default=str)

def load_history():
    items = []
    for f in sorted(HISTORY_DIR.glob("*.json"), reverse=True)[:50]:
        try:
            with open(f, "r", encoding="utf-8") as fh: items.append(json.load(fh))
        except: pass
    return items

# ---------------------------------------------------------------------------
# RENDER CARD
# ---------------------------------------------------------------------------
def render_patent_card(pat, show_relevance=True):
    app_html = f'<span class="applicant">{pat["applicant"]}</span>' if pat["applicant"] != "—" else ""
    link = f'<a href="{pat["gp_link"]}" target="_blank" style="font-size:0.75rem;color:#0ea5e9;text-decoration:none;">↗ Google Patents</a>' if pat.get("gp_link") else ""
    rel_html = ""
    if show_relevance and pat.get("relevance", 0) > 0:
        sc = pat["relevance"]
        cls = "rel-high" if sc >= 7 else ("rel-med" if sc >= 5 else "rel-low")
        rel_html = f'<span class="relevance {cls}">{sc}/10 · {pat.get("rel_type","—")}</span>'
        if pat.get("rel_reason","") not in ("","—","Nefiltrováno"):
            rel_html += f'<div class="meta" style="margin-top:4px;font-style:italic;">{pat["rel_reason"]}</div>'
    st.markdown(f"""<div class="patent-card">
        <h4>{pat['title']}</h4>
        <div class="meta">{pat['abstract'][:250]}{'...' if len(pat['abstract'])>250 else ''}</div>
        <div class="meta" style="margin-top:8px;"><strong>Č.:</strong> {pat['pub_number']} · <strong>Stát:</strong> {pat.get('country','—')} · <strong>Podáno:</strong> {pat['filing_date']} · {link}</div>
        {app_html} {rel_html}
    </div>""", unsafe_allow_html=True)


# ===========================================================================
# PAGE: HISTORY
# ===========================================================================
if st.session_state.page == "history":
    st.markdown("""<div class="hero"><div class="hero-badge">Archiv</div><h1>Historie analýz</h1><p>Přehled dosavadních rešerší</p></div>""", unsafe_allow_html=True)
    for item in load_history():
        with st.expander(f"📄 {item.get('pdf_name','—')} — {item.get('timestamp','—')}"):
            st.markdown(f"**Shrnutí:** {item.get('tech_summary','—')}")
            if item.get("analysis"): st.markdown(f'<div class="analysis-box">{item["analysis"][:2000]}...</div>', unsafe_allow_html=True)
    if not load_history(): st.info("Žádné uložené analýzy.")
    st.stop()

# ===========================================================================
# PAGE: MAIN
# ===========================================================================
st.markdown("""<div class="hero"><div class="hero-badge">Patent & Research Intelligence</div><h1>AI Transfer System</h1><p>Automatizovaná rešerše patentů a vědeckých publikací pro hodnocení komerčního potenciálu vynálezů</p><div class="hero-org">Fyzikální ústav AV ČR — Transfer znalostí a technologií</div></div>""", unsafe_allow_html=True)

phases = [("Upload PDF", 0), ("AI analýza", 1), ("Rešerše + filtr", 2), ("Statistiky", 3), ("Analýza", 4), ("Export", 5)]
phase_html = '<div class="phase-bar">'
for label, idx in phases:
    cls = "phase-done" if idx < st.session_state.phase else ("phase-active" if idx == st.session_state.phase else "phase-pending")
    phase_html += f'<span class="phase-pill {cls}">{label}</span>'
st.markdown(phase_html + "</div>", unsafe_allow_html=True)

# --- Upload ---
st.markdown("### 📄 Nahraj technický dokument")
uploaded = st.file_uploader("PDF patent, výzkumná zpráva, prezentace", type=["pdf"])
if uploaded and not st.session_state.pdf_text:
    with st.spinner("Extrahuji text..."): 
        text = extract_pdf_text(uploaded)
        if text: st.session_state.pdf_text = text; st.session_state.pdf_name = uploaded.name; st.session_state.phase = 1; st.rerun()
if st.session_state.pdf_text:
    with st.expander("📋 Extrahovaný text", expanded=False):
        st.text(st.session_state.pdf_text[:3000])

# --- Analyze doc ---
if st.session_state.phase >= 1 and not st.session_state.search_queries:
    st.markdown("---"); st.markdown("### 🧠 AI analýza dokumentu")
    if not gemini_key: st.warning("Zadej Gemini API klíč.")
    elif st.button("🚀 Analyzovat dokument", type="primary"):
        with st.spinner("Gemini analyzuje..."):
            try:
                r = analyze_document(gemini_key, st.session_state.pdf_text)
                st.session_state.tech_summary = r["summary"]; st.session_state.tech_keywords_en = r["keywords"]
                st.session_state.search_queries = r["queries"]
                if "domains" in r: st.session_state["tech_domains"] = r["domains"]
                st.session_state.phase = 2; st.rerun()
            except Exception as e: st.error(str(e))

if st.session_state.search_queries:
    st.markdown("---"); st.markdown("### 🔍 Vyhledávací dotazy")
    if st.session_state.tech_summary: st.info(f"**Shrnutí:** {st.session_state.tech_summary}")
    if st.session_state.get("tech_domains"): st.caption(f"**Identifikované aplikační domény:** {st.session_state.tech_domains}")
    if st.session_state.tech_keywords_en: st.caption(f"**Klíčová slova (EN):** {st.session_state.tech_keywords_en}")
    updated = []
    labels = ["🎯 Jádro technologie", "📱 Aplikace 1", "🏥 Aplikace 2", "⚡ Aplikace 3", "🧪 Materiál"]
    for i, q in enumerate(st.session_state.search_queries):
        lbl = labels[i] if i < len(labels) else f"Dotaz {i+1}"
        updated.append(st.text_input(lbl, value=q, key=f"q_{i}"))
    st.session_state.search_queries = updated

# --- Search + Filter ---
if st.session_state.phase >= 2 and st.session_state.search_queries and not st.session_state.patents_filtered:
    st.markdown("---")
    if not serpapi_key: st.warning("Zadej SerpApi Key.")
    if st.button("🔎 Spustit rešerši + AI filtr", type="primary", disabled=not serpapi_key):
        all_patents, seen = [], set()
        progress = st.progress(0, text="Rešerše...")
        for qi, q in enumerate(st.session_state.search_queries):
            sanitized = sanitize_patent_query(q)
            progress.progress(int((qi/len(st.session_state.search_queries))*35), text=f"Dotaz {qi+1}: {sanitized[:50]}...")
            try:
                results = search_google_patents(serpapi_key, q, max_patents_per_query)
                st.toast(f"Dotaz {qi+1}: nalezeno {len(results)} patentů")
                for p in results:
                    if p["pub_number"] not in seen: seen.add(p["pub_number"]); p["source_query"] = sanitized; all_patents.append(p)
            except Exception as e:
                st.error(f"Dotaz {qi+1} selhal: {e}\nDotaz: {sanitized}")
            time.sleep(0.3)

        st.session_state.patents_raw = all_patents

        if not all_patents:
            st.error("Nebyly nalezeny žádné patenty. Zkontroluj SerpApi klíč a dotazy.")
            progress.empty()
            st.stop()

        progress.progress(40, text=f"{len(all_patents)} patentů. AI filtr...")
        if gemini_key and all_patents:
            try:
                st.session_state.patents_filtered = filter_patents(gemini_key, all_patents, st.session_state.tech_summary, relevance_threshold)
            except Exception as e:
                st.warning(f"AI filtr selhal ({e}), zobrazuji všechny patenty")
                st.session_state.patents_filtered = all_patents
        else:
            st.session_state.patents_filtered = all_patents

        progress.progress(75, text="OpenAlex...")
        try:
            st.session_state.openalex_results = search_openalex(st.session_state.tech_keywords_en, max_openalex)
        except Exception as e:
            st.warning(f"OpenAlex selhal: {e}")

        progress.progress(100, text="Hotovo!")
        time.sleep(0.3); progress.empty()
        st.session_state.phase = 4; st.rerun()

# --- Results + Stats ---
if st.session_state.patents_raw:
    st.markdown("---"); st.markdown("### 📊 Výsledky rešerše")
    stats = compute_stats(st.session_state.patents_filtered)
    raw_c = len(st.session_state.patents_raw); filt_c = len(st.session_state.patents_filtered)
    competitors = [p for p in st.session_state.patents_filtered if "COMPET" in p.get("rel_type","").upper() or "KONKUR" in p.get("rel_type","").upper()]
    commercial = [r for r in st.session_state.openalex_results if r["is_commercial"]]

    st.markdown(f"""<div class="stat-row">
        <div class="stat-box"><div class="num">{raw_c} → {filt_c}</div><div class="label">Patentů</div></div>
        <div class="stat-box"><div class="num">{len(stats['assignees'])}</div><div class="label">Firem</div></div>
        <div class="stat-box"><div class="num">{len(st.session_state.openalex_results)}</div><div class="label">Publikací</div></div>
        <div class="stat-box"><div class="num">{len(commercial)}</div><div class="label">Komerční</div></div>
    </div>""", unsafe_allow_html=True)

    tab1, tab2, tab3, tab4 = st.tabs(["✅ Relevantní patenty", "📋 Všechny patenty", "📚 Publikace", "📈 Statistiky"])

    with tab1:
        for pat in st.session_state.patents_filtered: render_patent_card(pat)
    with tab2:
        st.caption(f"Všech {raw_c} patentů")
        for pat in st.session_state.patents_raw: render_patent_card(pat)
    with tab3:
        for pub in st.session_state.openalex_results:
            badge = "🏢 " if pub["is_commercial"] else ""
            st.markdown(f"""<div class="patent-card"><h4>{badge}{pub['title']}</h4>
                <div class="meta">Rok: {pub['year']} · Citací: {pub['cited_by']}× · Instituce: {', '.join(pub['institutions'][:3])}</div></div>""", unsafe_allow_html=True)
    with tab4:
        st.markdown("#### Top přihlašovatelé patentů")
        if stats["assignees"]:
            df_a = pd.DataFrame(stats["assignees"], columns=["Firma/Instituce", "Počet patentů"])
            st.dataframe(df_a, use_container_width=True, hide_index=True)
        st.markdown("#### Patentová aktivita po letech")
        if stats["yearly"]:
            df_y = pd.DataFrame(stats["yearly"], columns=["Rok", "Počet"])
            df_y = df_y[df_y["Počet"] > 0]
            st.bar_chart(df_y.set_index("Rok"), use_container_width=True)
        if stats["countries"]:
            st.markdown("#### Geografické rozložení")
            df_c = pd.DataFrame(stats["countries"], columns=["Stát", "Počet"])
            st.dataframe(df_c, use_container_width=True, hide_index=True)

# --- Analysis ---
if st.session_state.phase >= 4 and st.session_state.patents_filtered and not st.session_state.analysis:
    st.markdown("---"); st.markdown("### 🤖 Kritická AI analýza")
    if not gemini_key: st.warning("Zadej Gemini API klíč.")
    elif st.button("🧪 Spustit analýzu", type="primary"):
        with st.spinner("Gemini analyzuje..."):
            try:
                stats = compute_stats(st.session_state.patents_filtered)
                analysis = run_analysis(gemini_key, st.session_state.tech_summary, st.session_state.patents_filtered, st.session_state.openalex_results, st.session_state.pdf_text, stats)
                st.session_state.analysis = analysis; st.session_state.phase = 5
                save_to_history({"timestamp": datetime.now().strftime("%d.%m.%Y %H:%M"), "pdf_name": st.session_state.pdf_name,
                    "tech_summary": st.session_state.tech_summary, "search_queries": st.session_state.search_queries,
                    "patents_raw_count": len(st.session_state.patents_raw),
                    "patents_filtered": [{"title":p["title"],"applicant":p["applicant"],"pub_number":p["pub_number"],"relevance":p["relevance"],"rel_type":p["rel_type"]} for p in st.session_state.patents_filtered],
                    "openalex_results": [{"title":p["title"],"year":p["year"],"cited_by":p["cited_by"],"is_commercial":p["is_commercial"],"institutions":p["institutions"]} for p in st.session_state.openalex_results],
                    "analysis": analysis})
                st.rerun()
            except Exception as e: st.error(str(e))

if st.session_state.analysis:
    st.markdown("---"); st.markdown("### 🤖 Výsledek analýzy")
    st.markdown(f'<div class="analysis-box">{st.session_state.analysis}</div>', unsafe_allow_html=True)

# --- Export ---
if st.session_state.phase >= 5 and st.session_state.analysis:
    st.markdown("---"); st.markdown("### 📥 Export")
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("📄 Word report", type="primary"):
            with st.spinner("Generuji .docx..."):
                try:
                    stats = compute_stats(st.session_state.patents_filtered)
                    buf = generate_docx(st.session_state.tech_summary, st.session_state.search_queries,
                        len(st.session_state.patents_raw), st.session_state.patents_filtered,
                        st.session_state.openalex_results, st.session_state.analysis,
                        st.session_state.pdf_name, relevance_threshold, stats)
                    st.session_state.doc_content = buf.getvalue(); st.success("Hotovo!")
                except Exception as e: st.error(str(e))
    with col2:
        if st.button("📊 Excel tabulka", type="primary"):
            with st.spinner("Generuji .xlsx..."):
                try:
                    buf = generate_xlsx(st.session_state.patents_filtered, st.session_state.openalex_results, st.session_state.tech_summary)
                    st.session_state.xlsx_content = buf.getvalue(); st.success("Hotovo!")
                except Exception as e: st.error(str(e))
    with col3:
        pass
    dl1, dl2, dl3 = st.columns(3)
    if st.session_state.doc_content:
        with dl1: st.download_button("⬇️ Stáhnout .docx", data=st.session_state.doc_content,
            file_name=f"reserse_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if st.session_state.xlsx_content:
        with dl2: st.download_button("⬇️ Stáhnout .xlsx", data=st.session_state.xlsx_content,
            file_name=f"patenty_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

if st.session_state.phase > 0:
    st.markdown("---")
    if st.button("🔄 Nová analýza (reset)"):
        for k in defaults: st.session_state[k] = defaults[k]
        st.rerun()
